from __future__ import annotations

import os
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str((Path.cwd() / ".mplconfig").resolve()))
Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches
from matplotlib import font_manager



def setup_plotting() -> None:
    font_candidates = [Path(r"C:\Windows\Fonts\msyh.ttc"), Path(r"C:\Windows\Fonts\simhei.ttf")]
    font_name = "DejaVu Sans"
    for font_path in font_candidates:
        if font_path.exists():
            font_manager.fontManager.addfont(str(font_path))
            font_name = font_manager.FontProperties(fname=str(font_path)).get_name()
            break
    sns.set_theme(style="whitegrid")
    plt.rcParams["font.family"] = font_name
    plt.rcParams["font.sans-serif"] = [font_name, "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False



def save_distribution_plots(frame: pd.DataFrame, output_dir: Path) -> pd.DataFrame:
    numeric_cols = [c for c in frame.columns if c != "日期" and pd.api.types.is_numeric_dtype(frame[c])]
    stats = frame[numeric_cols].describe().T.reset_index().rename(columns={"index": "字段"})
    page_size = 6
    for start in range(0, len(numeric_cols), page_size):
        cols = numeric_cols[start : start + page_size]
        fig, axes = plt.subplots(len(cols), 2, figsize=(12, max(4, len(cols) * 3)))
        if len(cols) == 1:
            axes = [axes]
        for row_idx, col in enumerate(cols):
            sns.histplot(frame[col].dropna(), kde=True, ax=axes[row_idx][0], color="#4C72B0")
            axes[row_idx][0].set_title(f"{col} 直方图")
            sns.boxplot(x=frame[col], ax=axes[row_idx][1], color="#55A868")
            axes[row_idx][1].set_title(f"{col} 箱线图")
        fig.tight_layout()
        fig.savefig(output_dir / f"distribution_page_{start // page_size + 1}.png", dpi=200, bbox_inches="tight")
        plt.close(fig)
    return stats



def add_period_shading(ax, years: list[int]) -> None:
    for year in years:
        ax.axvspan(pd.Timestamp(year=year, month=4, day=1), pd.Timestamp(year=year, month=6, day=30), color="#F6C85F", alpha=0.12)
        ax.axvspan(pd.Timestamp(year=year, month=6, day=1), pd.Timestamp(year=year, month=10, day=31), color="#6F4E7C", alpha=0.08)



def save_trend_plots(frame: pd.DataFrame, output_dir: Path) -> None:
    targets = ["原水浊度均值", "原水pH均值", "原水温度均值", "对齐后投矾目标"]
    existing = [c for c in targets if c in frame.columns]
    fig, axes = plt.subplots(len(existing), 1, figsize=(15, max(4, len(existing) * 3)), sharex=True)
    if len(existing) == 1:
        axes = [axes]
    years = sorted(frame["日期"].dt.year.unique().tolist())
    for ax, col in zip(axes, existing):
        ax.plot(frame["日期"], frame[col], linewidth=1.1)
        ax.set_title(f"{col} 时间趋势")
        add_period_shading(ax, years)
    fig.tight_layout()
    fig.savefig(output_dir / "time_series_trends.png", dpi=200, bbox_inches="tight")
    plt.close(fig)



def save_correlation_heatmap(frame: pd.DataFrame, output_dir: Path, target_col: str) -> pd.DataFrame:
    numeric = frame.select_dtypes(include=["number"]).copy()
    corr = numeric.corr(numeric_only=True)
    focus = corr[[target_col]].sort_values(target_col, ascending=False)
    top_cols = focus[target_col].abs().sort_values(ascending=False).head(min(18, len(focus))).index.tolist()
    top_corr = corr.loc[top_cols, top_cols]
    plt.figure(figsize=(12, 10))
    sns.heatmap(top_corr, cmap="RdBu_r", center=0, annot=False)
    plt.title("关键特征相关性热力图")
    plt.tight_layout()
    plt.savefig(output_dir / "correlation_heatmap.png", dpi=220, bbox_inches="tight")
    plt.close()
    return focus.reset_index().rename(columns={"index": "字段", target_col: "与目标相关系数"})



def save_operating_condition_table(frame: pd.DataFrame, output_dir: Path, target_col: str) -> pd.DataFrame:
    rows = []
    definitions = {
        "低浊期": frame["工况_低浊期"] == 1,
        "汛期": frame["工况_汛期"] == 1,
        "低温期": frame["工况_低温期"] == 1,
        "分层期": frame["工况_分层期"] == 1,
    }
    for name, mask in definitions.items():
        subset = frame[mask]
        if subset.empty:
            continue
        rows.append(
            {
                "工况": name,
                "样本数": len(subset),
                "平均投矾": subset[target_col].mean(),
                "投矾标准差": subset[target_col].std(),
                "平均浊度": subset.get("原水浊度均值", pd.Series(dtype=float)).mean(),
                "平均pH": subset.get("原水pH均值", pd.Series(dtype=float)).mean(),
                "平均温度": subset.get("原水温度均值", pd.Series(dtype=float)).mean(),
            }
        )
    result = pd.DataFrame(rows)
    result.to_excel(output_dir / "operating_conditions.xlsx", index=False)
    return result



def save_model_plots(predictions: pd.DataFrame, importance: pd.DataFrame, convergence: list[float], output_dir: Path) -> None:
    pred_col = "推荐模型预测"
    plt.figure(figsize=(6, 6))
    sns.scatterplot(data=predictions, x="真实值", y=pred_col, s=25)
    low = min(predictions["真实值"].min(), predictions[pred_col].min())
    high = max(predictions["真实值"].max(), predictions[pred_col].max())
    plt.plot([low, high], [low, high], linestyle="--", color="red")
    plt.title("预测值 vs 真实值")
    plt.tight_layout()
    plt.savefig(output_dir / "prediction_vs_actual_scatter.png", dpi=220, bbox_inches="tight")
    plt.close()

    show_n = min(100, len(predictions))
    subset = predictions.iloc[:show_n]
    plt.figure(figsize=(14, 5))
    plt.plot(subset["日期"], subset["真实值"], label="真实值", linewidth=1.4)
    plt.plot(subset["日期"], subset[pred_col], label="预测值", linewidth=1.2)
    plt.legend()
    plt.title("连续100点时序预测对比")
    plt.tight_layout()
    plt.savefig(output_dir / "time_series_prediction_100_points.png", dpi=220, bbox_inches="tight")
    plt.close()

    plt.figure(figsize=(8, 5))
    sns.histplot(predictions["误差"], bins=30, kde=True, color="#C44E52")
    plt.title("预测误差分布")
    plt.tight_layout()
    plt.savefig(output_dir / "error_distribution.png", dpi=220, bbox_inches="tight")
    plt.close()

    importance_col = "\u878d\u5408\u7ebf\u6027\u7cfb\u6570\u7edd\u5bf9\u503c" if "\u878d\u5408\u7ebf\u6027\u7cfb\u6570\u7edd\u5bf9\u503c" in importance.columns else ("ElasticNet\u7cfb\u6570\u7edd\u5bf9\u503c" if "ElasticNet\u7cfb\u6570\u7edd\u5bf9\u503c" in importance.columns else importance.columns[-1])
    top_importance = importance.head(20).sort_values(importance_col)
    plt.figure(figsize=(10, 8))
    plt.barh(top_importance["特征"], top_importance[importance_col], color="#4C72B0")
    plt.title("特征重要性 Top 20")
    plt.tight_layout()
    plt.savefig(output_dir / "feature_importance_top20.png", dpi=220, bbox_inches="tight")
    plt.close()

    plt.figure(figsize=(8, 4.5))
    plt.plot(range(len(convergence)), convergence, marker="o", markersize=2, linewidth=1.2)
    plt.title("ISSA 收敛曲线")
    plt.xlabel("迭代次数")
    plt.ylabel("验证集 RMSE")
    plt.tight_layout()
    plt.savefig(output_dir / "issa_convergence.png", dpi=220, bbox_inches="tight")
    plt.close()



def build_word_report(
    output_path: Path,
    summary_text: list[str],
    metrics: pd.DataFrame,
    comparison: pd.DataFrame,
    stability: pd.DataFrame,
    generalization: pd.DataFrame,
    image_paths: list[Path],
) -> None:
    doc = Document()
    doc.add_heading("混凝智能投药预测模型技术报告", level=0)
    for paragraph in summary_text:
        doc.add_paragraph(paragraph)

    sections = [
        ("测试集评估指标", metrics),
        ("对比实验结果", comparison),
        ("稳定性测试", stability),
        ("工况泛化测试", generalization),
    ]
    for title, frame in sections:
        doc.add_heading(title, level=1)
        if frame.empty:
            doc.add_paragraph("当前无可用结果。")
            continue
        table = doc.add_table(rows=1, cols=len(frame.columns))
        for idx, col in enumerate(frame.columns):
            table.rows[0].cells[idx].text = str(col)
        for _, row in frame.iterrows():
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                if isinstance(value, float):
                    cells[idx].text = f"{value:.4f}"
                else:
                    cells[idx].text = str(value)

    doc.add_heading("关键图表", level=1)
    for image in image_paths:
        if image.exists():
            doc.add_paragraph(image.stem)
            doc.add_picture(str(image), width=Inches(6.5))
    doc.save(output_path)
